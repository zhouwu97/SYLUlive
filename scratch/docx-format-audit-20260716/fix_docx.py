"""按竞赛提交规范修正作品说明文档格式。"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


FONT_NAME = "宋体"
BODY_SIZE = 10.5


def set_rfonts(element, font_name: str = FONT_NAME):
    """同时设置中西文字体，避免 Word/WPS 按主题字体回退。"""
    rpr = element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font_name)
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        key = qn(f"w:{attr}")
        if key in rfonts.attrib:
            del rfonts.attrib[key]


def format_run(run, size_pt: float, bold: bool | None = None):
    run.font.name = FONT_NAME
    set_rfonts(run._element)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def format_style(style, size_pt: float, bold: bool | None = None):
    style.font.name = FONT_NAME
    style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    set_rfonts(style._element)
    if style.type == WD_STYLE_TYPE.PARAGRAPH:
        style.paragraph_format.line_spacing = 1.0


def ensure_paragraph_style(doc, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def configure_styles(doc):
    """把模板样式和标准标题样式都固定到竞赛要求的字号。"""
    style_specs = {
        "Normal": (BODY_SIZE, False),
        "Normal (Web)": (BODY_SIZE, False),
        "正文": (BODY_SIZE, False),
        "申报正文": (BODY_SIZE, False),
        "Title": (22, True),
        "申报标题": (22, True),
        "Heading 1": (16, True),
        "申报一级标题": (16, True),
        "Heading 2": (14, True),
        "申报二级标题": (14, True),
        "Heading 3": (12, True),
        "申报三级标题": (12, True),
        "Heading 4": (10.5, True),
        "申报四级标题": (10.5, True),
        "Caption": (BODY_SIZE, False),
        "申报图注": (BODY_SIZE, False),
        "Footer": (9, False),
    }
    for name, (size, bold) in style_specs.items():
        style = ensure_paragraph_style(doc, name)
        format_style(style, size, bold)


def insert_paragraph_before_table(table, text: str, style_name: str):
    new_element = OxmlElement("w:p")
    table._tbl.addprevious(new_element)
    paragraph = Paragraph(new_element, table._parent)
    paragraph.style = style_name
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    format_run(run, BODY_SIZE, False)
    return paragraph


def set_row_nonbreaking(row):
    trpr = row._tr.get_or_add_trPr()
    if trpr.find(qn("w:cantSplit")) is None:
        trpr.append(OxmlElement("w:cantSplit"))
    # 固定行高可能截断文字，改为最小行高并允许自然扩展。
    for height in trpr.findall(qn("w:trHeight")):
        if height.get(qn("w:hRule")) == "exact":
            height.set(qn("w:hRule"), "atLeast")


def iter_unique_cells(row):
    seen = set()
    for cell in row.cells:
        key = id(cell._tc)
        if key not in seen:
            seen.add(key)
            yield cell


def format_table(table):
    for row in table.rows:
        set_row_nonbreaking(row)
        for cell in iter_unique_cells(row):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    if run.text:
                        format_run(run, BODY_SIZE)


def add_page_field(paragraph):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.extend([begin, instr, separate, value, end])
    format_run(run, 9, False)


def clear_header_footer(part):
    for child in list(part._element):
        part._element.remove(child)


def footer_text(paragraph, text: str):
    run = paragraph.add_run(text)
    format_run(run, 9, False)


def build_footer(footer, even_page: bool):
    """镜像页脚：偶数页页码靠左，奇数页页码靠右。"""
    clear_header_footer(footer)
    paragraph = footer.add_paragraph(style="Footer")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(7.5), WD_TAB_ALIGNMENT.CENTER)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT)

    committee = "中国高校计算机大赛-人工智能创意赛（鸿蒙赛道）组委会编制"
    version = "2026.06"
    if even_page:
        footer_text(paragraph, "第 ")
        add_page_field(paragraph)
        footer_text(paragraph, " 页\t")
        footer_text(paragraph, version)
        footer_text(paragraph, "\t" + committee)
    else:
        footer_text(paragraph, committee)
        footer_text(paragraph, "\t" + version + "\t第 ")
        add_page_field(paragraph)
        footer_text(paragraph, " 页")


def configure_page_and_footers(doc):
    settings = doc.settings.element
    if settings.find(qn("w:evenAndOddHeaders")) is None:
        settings.append(OxmlElement("w:evenAndOddHeaders"))

    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
        section.gutter = Cm(0)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)
        section.different_first_page_header_footer = False
        section.footer.is_linked_to_previous = False
        section.even_page_footer.is_linked_to_previous = False
        build_footer(section.footer, even_page=False)
        build_footer(section.even_page_footer, even_page=True)


def replace_exact_text(paragraph, replacements):
    text = paragraph.text.strip()
    if text in replacements:
        paragraph.text = replacements[text]
        return replacements[text]
    return text


def format_body_paragraphs(doc):
    replacements = {
        "参赛团队信息表": "表 1 参赛团队信息表",
        "创意描述": "一、创意描述",
        "三、技术方案与 AI 处理流程": "四、技术方案与 AI 处理流程",
        "四、作品介绍文档": "五、作品介绍文档",
        "图1 校园社区信息流            图2 校园资讯与服务聚合": (
            "图 1 校园社区信息流            图 2 校园资讯与服务聚合"
        ),
    }
    title_texts = {
        "2026中国高校计算机大赛—人工智能创意赛",
        "鸿蒙赛道作品说明文档（初赛）",
        "作品说明文档（初赛）",
        "表 1 参赛团队信息表",
        "《沈理校园（SYLUlive）》作品原创性声明",
        "沈理校园（SYLUlive）作品说明",
    }
    author_prefixes = (
        "参赛学校：",
        "团队名称：",
        "作品名称：",
        "赛题方向：",
        "联系人（队长）：",
        "联系电话（队长）：",
    )

    for paragraph in doc.paragraphs:
        text = replace_exact_text(paragraph, replacements)
        paragraph.paragraph_format.line_spacing = 1.0

        style_name = paragraph.style.name if paragraph.style else ""
        if text in title_texts:
            size, bold = 22, True
            paragraph.paragraph_format.keep_with_next = True
        elif text.startswith(author_prefixes):
            size, bold = 16, True
        elif text == "包含但不限于以下内容：":
            size, bold = 16, True
            paragraph.paragraph_format.keep_with_next = True
        elif style_name == "申报一级标题" or text.startswith(
            ("一、", "二、", "三、", "四、", "五、")
        ):
            size, bold = 16, True
            paragraph.paragraph_format.keep_with_next = True
        elif style_name == "申报二级标题":
            size, bold = 14, True
            paragraph.paragraph_format.keep_with_next = True
        elif style_name == "申报三级标题":
            size, bold = 12, True
            paragraph.paragraph_format.keep_with_next = True
        elif style_name == "申报四级标题":
            size, bold = 10.5, True
            paragraph.paragraph_format.keep_with_next = True
        elif style_name in {"申报图注", "Caption"} or text.startswith(("图 ", "表 ")):
            size, bold = BODY_SIZE, False
        else:
            size, bold = BODY_SIZE, None

        for run in paragraph.runs:
            if run.text:
                format_run(run, size, bold)

        if text.startswith("状态说明："):
            paragraph.paragraph_format.space_before = Pt(4)


def main():
    if len(sys.argv) != 3:
        raise SystemExit("用法: fix_docx.py 输入.docx 输出.docx")

    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(source)
    configure_styles(doc)
    format_body_paragraphs(doc)

    for table in doc.tables:
        format_table(table)
    if len(doc.tables) >= 2:
        insert_paragraph_before_table(
            doc.tables[1], "表 2 实现基础与鸿蒙适配计划", "申报图注"
        )

    configure_page_and_footers(doc)
    doc.core_properties.title = "沈理校园（SYLUlive）作品说明文档（初赛）"
    doc.core_properties.subject = "2026中国高校计算机大赛人工智能创意赛鸿蒙赛道"
    doc.core_properties.author = "咕咕嘎嘎"
    doc.core_properties.last_modified_by = "咕咕嘎嘎"
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
