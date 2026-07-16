"""审计作品说明文档的版式、样式、题注和媒体资源。"""

from __future__ import annotations

import json
import re
import sys
import zipfile
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image


def length_pt(value):
    return None if value is None else round(value.pt, 2)


def font_name(run):
    rpr = run._element.rPr
    east_asia = None
    if rpr is not None and rpr.rFonts is not None:
        east_asia = rpr.rFonts.get(qn("w:eastAsia"))
    return east_asia or run.font.name


def paragraph_kind(text: str) -> str:
    clean = text.strip()
    if not clean:
        return "空段落"
    if re.match(r"^(参考资料|参考文献|附录)($|\s|[一二三四五六七八九十0-9A-Z])", clean):
        return "参考或附录"
    if re.match(r"^第[一二三四五六七八九十0-9]+章", clean):
        return "一级标题候选"
    if re.match(r"^[一二三四五六七八九十]+、", clean):
        return "一级标题候选"
    if re.match(r"^（[一二三四五六七八九十]+）", clean):
        return "二级标题候选"
    if re.match(r"^[0-9]+[、.]", clean):
        return "三级标题候选"
    if re.match(r"^（[0-9]+）", clean):
        return "四级标题候选"
    if re.match(r"^(图|表)\s*[0-9一二三四五六七八九十]+", clean):
        return "题注候选"
    return "正文候选"


def run_signature(run):
    return {
        "font": font_name(run),
        "size_pt": length_pt(run.font.size),
        "bold": run.bold,
        "italic": run.italic,
        "text": run.text[:80],
    }


def effective_paragraph_signature(paragraph):
    pf = paragraph.paragraph_format
    style = paragraph.style
    style_pf = style.paragraph_format if style is not None else None
    line_spacing = pf.line_spacing
    if line_spacing is None and style_pf is not None:
        line_spacing = style_pf.line_spacing
    return {
        "style": style.name if style is not None else None,
        "alignment": str(paragraph.alignment),
        "line_spacing": str(line_spacing),
        "space_before_pt": length_pt(pf.space_before),
        "space_after_pt": length_pt(pf.space_after),
        "first_line_indent_pt": length_pt(pf.first_line_indent),
    }


def audit(path: Path):
    doc = Document(path)
    paragraph_rows = []
    style_counts = Counter()
    kind_counts = Counter()
    run_font_counts = Counter()
    run_size_counts = Counter()
    run_bold_counts = Counter()

    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        kind = paragraph_kind(text)
        style_counts[paragraph.style.name if paragraph.style else "(无样式)"] += 1
        kind_counts[kind] += 1
        run_rows = []
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            sig = run_signature(run)
            run_rows.append(sig)
            run_font_counts[str(sig["font"])] += len(run.text.strip())
            run_size_counts[str(sig["size_pt"])] += len(run.text.strip())
            run_bold_counts[str(sig["bold"])] += len(run.text.strip())
        paragraph_rows.append(
            {
                "index": index,
                "kind": kind,
                "text": text[:180],
                "paragraph": effective_paragraph_signature(paragraph),
                "runs": run_rows,
            }
        )

    table_rows = []
    for table_index, table in enumerate(doc.tables, start=1):
        cells = []
        for row_index, row in enumerate(table.rows, start=1):
            for col_index, cell in enumerate(row.cells, start=1):
                text = "\n".join(p.text for p in cell.paragraphs).strip()
                cells.append(
                    {
                        "row": row_index,
                        "column": col_index,
                        "text": text[:180],
                        "paragraphs": [
                            {
                                "text": p.text[:120],
                                "style": p.style.name if p.style else None,
                                "runs": [run_signature(r) for r in p.runs if r.text.strip()],
                            }
                            for p in cell.paragraphs
                        ],
                    }
                )
        table_rows.append(
            {
                "index": table_index,
                "style": table.style.name if table.style else None,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "cells": cells,
            }
        )

    media_rows = []
    with zipfile.ZipFile(path) as package:
        for name in sorted(n for n in package.namelist() if n.startswith("word/media/")):
            data = package.read(name)
            media = {"name": name, "bytes": len(data), "valid": False}
            try:
                from io import BytesIO

                with Image.open(BytesIO(data)) as image:
                    image.verify()
                    media.update(
                        {
                            "valid": True,
                            "format": image.format,
                            "width": image.width,
                            "height": image.height,
                        }
                    )
            except Exception as exc:  # noqa: BLE001
                media["error"] = repr(exc)
            media_rows.append(media)

    sections = []
    for index, section in enumerate(doc.sections, start=1):
        sections.append(
            {
                "index": index,
                "page_width_cm": round(section.page_width.cm, 3),
                "page_height_cm": round(section.page_height.cm, 3),
                "top_margin_cm": round(section.top_margin.cm, 3),
                "bottom_margin_cm": round(section.bottom_margin.cm, 3),
                "left_margin_cm": round(section.left_margin.cm, 3),
                "right_margin_cm": round(section.right_margin.cm, 3),
                "header_distance_cm": round(section.header_distance.cm, 3),
                "footer_distance_cm": round(section.footer_distance.cm, 3),
                "gutter_cm": round(section.gutter.cm, 3),
            }
        )

    core = doc.core_properties
    return {
        "path": str(path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "inline_shape_count": len(doc.inline_shapes),
        "section_count": len(doc.sections),
        "sections": sections,
        "core_properties": {
            "title": core.title,
            "author": core.author,
            "last_modified_by": core.last_modified_by,
            "revision": core.revision,
        },
        "style_counts": dict(style_counts.most_common()),
        "kind_counts": dict(kind_counts.most_common()),
        "run_font_weighted_counts": dict(run_font_counts.most_common()),
        "run_size_weighted_counts": dict(run_size_counts.most_common()),
        "run_bold_weighted_counts": dict(run_bold_counts.most_common()),
        "paragraphs": paragraph_rows,
        "tables": table_rows,
        "media": media_rows,
    }


def main():
    if len(sys.argv) != 3:
        raise SystemExit("用法: audit_docx.py 输入.docx 输出.json")
    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    result = audit(source)
    output.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    summary = {key: value for key, value in result.items() if key not in {"paragraphs", "tables"}}
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
