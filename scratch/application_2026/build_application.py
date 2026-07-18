from __future__ import annotations

from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\python_play_do\sylg-live")
WORK = ROOT / "scratch" / "application_2026"
TEMPLATE = WORK / "template.docx"
OUTPUT = ROOT / "01-作品说明文档+咕咕嘎嘎.docx"

SCREENSHOTS = [
    Path(r"C:\Users\zhy23\AppData\Local\Temp\codex-clipboard-4fec1312-9962-421e-9e23-941ca365e8bd.png"),
    Path(r"C:\Users\zhy23\AppData\Local\Temp\codex-clipboard-4531ecc4-757b-43f5-8d03-540a008c0c73.png"),
    Path(r"C:\Users\zhy23\AppData\Local\Temp\codex-clipboard-048161d9-0148-446f-b2d0-f7fd40c571d2.png"),
]

PRESERVE_ONLY_PARTS = {
    "customXml/_rels/item1.xml.rels",
    "word/numbering.xml",
}

FONT_CN = "宋体"
FONT_LATIN = "Times New Roman"
COLOR_TEXT = RGBColor(31, 35, 41)
COLOR_MUTED = RGBColor(96, 104, 116)
COLOR_ACCENT = RGBColor(0, 119, 105)
COLOR_ACCENT_DARK = RGBColor(0, 88, 78)


def set_run_font(run, size: float, *, bold: bool = False, color: RGBColor | None = None) -> None:
    """统一设置中西文字体，避免不同渲染器发生中文字体回退。"""
    run.font.name = FONT_LATIN
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_fonts.set(qn("w:eastAsia"), FONT_CN)


def set_cell_margins(cell, *, top: int = 100, start: int = 100, bottom: int = 100, end: int = 100) -> None:
    """以 DXA 设置单元格内边距，保证表格文字不贴边。"""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm: float) -> None:
    width_dxa = int(Cm(width_cm).twips)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Cm(width_cm)


def mark_header_row(row) -> None:
    """标记 Word 表头行，提升屏幕阅读器识别和跨页重复能力。"""
    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")


def set_table_geometry(table, widths_cm: list[float]) -> None:
    """同步设置 tblW、tblGrid 与 tcW，避免自动列宽在不同软件中漂移。"""
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total_dxa = sum(int(Cm(width).twips) for width in widths_cm)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(Cm(width).twips)))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_cm[index])


def remove_template_instructions(doc: Document) -> None:
    """移除第 5、6 页填写规范，保留第 4 页末尾的分页符。"""
    body = doc.element.body
    start = None
    for child in body.iterchildren():
        if child.tag != qn("w:p"):
            continue
        text = "".join(child.itertext())
        if "作品说明文档提交规范说明" in text:
            start = child
            break
    if start is None:
        raise RuntimeError("未找到模板规范说明起点，拒绝继续生成。")

    remove = False
    for child in list(body):
        if child is start:
            remove = True
        if remove and child.tag != qn("w:sectPr"):
            body.remove(child)


def fix_front_matter(doc: Document) -> None:
    """只修正已填内容中的明显病句和标点，不改动报名事实。"""
    if not doc.tables:
        raise RuntimeError("模板团队信息表缺失。")
    mark_header_row(doc.tables[0].rows[0])
    cell = doc.tables[0].rows[13].cells[0]
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if "还要多个项目经历" in run.text:
                run.text = run.text.replace("还要多个项目经历", "拥有多项项目经历")

    for paragraph in doc.paragraphs:
        if "作品原创性声明" in paragraph.text and "沈理校园" in paragraph.text:
            paragraph.text = "《沈理校园（SYLUlive）》作品原创性声明"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, 18, bold=True)
            break


def create_styles(doc: Document) -> None:
    styles = doc.styles

    def ensure(name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
        return styles[name] if name in styles else styles.add_style(name, style_type)

    body = ensure("申报正文")
    body.font.name = FONT_LATIN
    body.font.size = Pt(10.5)
    body._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    body.paragraph_format.line_spacing = 1.0
    body.paragraph_format.space_after = Pt(3)
    body.paragraph_format.first_line_indent = Pt(21)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title = ensure("申报标题")
    title.font.name = FONT_LATIN
    title.font.size = Pt(22)
    title.font.bold = True
    title._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.keep_with_next = True

    h1 = ensure("申报一级标题")
    h1.font.name = FONT_LATIN
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    h1.paragraph_format.space_before = Pt(7)
    h1.paragraph_format.space_after = Pt(5)
    h1.paragraph_format.keep_with_next = True

    h2 = ensure("申报二级标题")
    h2.font.name = FONT_LATIN
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    h2.paragraph_format.space_before = Pt(5)
    h2.paragraph_format.space_after = Pt(3)
    h2.paragraph_format.keep_with_next = True

    caption = ensure("申报图注")
    caption.font.name = FONT_LATIN
    caption.font.size = Pt(9)
    caption.font.color.rgb = COLOR_MUTED
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.keep_with_previous = True


def add_heading(doc: Document, text: str, level: int = 1):
    style = "申报一级标题" if level == 1 else "申报二级标题"
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, 16 if level == 1 else 14, bold=True, color=COLOR_ACCENT_DARK if level == 1 else COLOR_TEXT)
    return p


def add_body(doc: Document, text: str, *, indent: bool = True):
    p = doc.add_paragraph(style="申报正文")
    if not indent:
        p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, 10.5, color=COLOR_TEXT)
    return p


def add_labeled_body(doc: Document, label: str, text: str):
    p = doc.add_paragraph(style="申报正文")
    p.paragraph_format.first_line_indent = Pt(0)
    r1 = p.add_run(label)
    set_run_font(r1, 10.5, bold=True, color=COLOR_ACCENT_DARK)
    r2 = p.add_run(text)
    set_run_font(r2, 10.5, color=COLOR_TEXT)
    return p


def add_callout(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "E8F4F1")
    p_pr.append(shd)
    r = p.add_run(text)
    set_run_font(r, 14, bold=True, color=COLOR_ACCENT_DARK)
    return p


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_screenshot_pair(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    image1 = p.add_run().add_picture(str(SCREENSHOTS[0]), width=Cm(4.8))
    image1._inline.docPr.set("title", "图1 校园社区信息流")
    image1._inline.docPr.set("descr", "沈理校园社区首页，展示帖子、图片以及浏览、点赞和评论数据。")
    p.add_run("   ")
    image2 = p.add_run().add_picture(str(SCREENSHOTS[1]), width=Cm(4.8))
    image2._inline.docPr.set("title", "图2 校园资讯与服务聚合")
    image2._inline.docPr.set("descr", "沈理校园校园页，展示竞赛通知、校园资讯和教务、组队、地图、校历入口。")

    cap = doc.add_paragraph(style="申报图注")
    cap.paragraph_format.tab_stops.add_tab_stop(Cm(3.65))
    cap.paragraph_format.tab_stops.add_tab_stop(Cm(10.95))
    run = cap.add_run("\t图1 校园社区信息流\t图2 校园资讯与服务聚合")
    set_run_font(run, 9, color=COLOR_MUTED)


def add_single_screenshot(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    image3 = p.add_run().add_picture(str(SCREENSHOTS[2]), width=Cm(7.0))
    image3._inline.docPr.set("title", "图3 校园服务入口与功能分区")
    image3._inline.docPr.set("descr", "沈理校园服务侧边面板，按公告、社区板块和竞赛、成绩、考试、反馈等任务分区。")
    cap = doc.add_paragraph("图3 校园服务入口与功能分区", style="申报图注")
    for run in cap.runs:
        set_run_font(run, 9, color=COLOR_MUTED)


def add_technical_table(doc: Document) -> None:
    rows = [
        ("客户端界面", "Flutter/Android 原型已可运行，完成社区首页、校园资讯和服务入口。", "使用 ArkTS、ArkUI 与 Stage 模型重构核心页面，计划适配服务卡片和系统通知。"),
        ("业务服务", "Go REST API 承担 JWT 鉴权、社区、消息、评价及业务编排。", "复用既有 JSON 接口，增加鸿蒙设备订阅、提醒状态和跨端任务同步。"),
        ("数据与智能", "Python 服务承担教务及校园公开信息采集，接口支持竞赛、成绩等数据。", "在既有链路上新增大模型分类、摘要和字段抽取，并执行规则校验、去重、来源回链。"),
        ("存储与部署", "PostgreSQL 持久化，Nginx 反向代理，Docker 统一部署。", "采用数据最小化、敏感字段脱敏、日志审计和失败回退，保障校内落地稳定性。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    mark_header_row(table.rows[0])
    widths = [2.6, 5.8, 6.2]
    hdr = table.rows[0].cells
    headers = ["层级", "已完成基础", "AI 增强与鸿蒙赛事版计划"]
    for idx, text in enumerate(headers):
        hdr[idx].text = text
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(hdr[idx], "DDEEEA")
        set_cell_margins(hdr[idx], top=120, start=110, bottom=120, end=110)
        for p in hdr[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_run_font(run, 9.5, bold=True, color=COLOR_ACCENT_DARK)

    for row_data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            cells[idx].text = text
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[idx], top=110, start=110, bottom=110, end=110)
            for p in cells[idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, 9, bold=(idx == 0), color=COLOR_TEXT)
    set_table_geometry(table, widths)


def add_page_number(doc: Document) -> None:
    """将动态页码并入既有页脚同一行，避免挤动模板正文。"""
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.clear_all()
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(14.55), WD_TAB_ALIGNMENT.RIGHT)
    left = paragraph.add_run("中国高校计算机大赛-人工智能创意赛（鸿蒙赛道）组委会编制")
    set_run_font(left, 9, color=COLOR_TEXT)
    run = paragraph.add_run("\t2026.06  第 ")
    set_run_font(run, 9, color=COLOR_MUTED)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    cached = OxmlElement("w:t")
    cached.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, cached, end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, 9, color=COLOR_MUTED)

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def restore_preserve_only_parts() -> None:
    """恢复无需编辑的模板部件，避免库保存时产生无意义的 XML 重排。"""
    patched = OUTPUT.with_suffix(".preserved.docx")
    with ZipFile(TEMPLATE, "r") as reference, ZipFile(OUTPUT, "r") as generated, ZipFile(
        patched, "w", compression=ZIP_DEFLATED
    ) as result:
        for info in generated.infolist():
            data = reference.read(info.filename) if info.filename in PRESERVE_ONLY_PARTS else generated.read(info.filename)
            result.writestr(info, data)
    patched.replace(OUTPUT)


def build() -> None:
    for path in [TEMPLATE, *SCREENSHOTS]:
        if not path.exists():
            raise FileNotFoundError(path)

    doc = Document(TEMPLATE)
    fix_front_matter(doc)
    remove_template_instructions(doc)
    create_styles(doc)
    add_page_number(doc)

    title = doc.add_paragraph("沈理校园（SYLUlive）作品说明", style="申报标题")
    for run in title.runs:
        set_run_font(run, 22, bold=True, color=COLOR_TEXT)

    add_heading(doc, "一、创意描述")
    add_callout(doc, "让校园信息在鸿蒙端智能聚合、精准触达")

    add_heading(doc, "二、产品定位与创新价值")
    add_labeled_body(doc, "真实痛点：", "教务、学院、竞赛和学生社区信息分散，学生需要跨网站、群聊反复查找，重要事项容易遗漏。")
    add_labeled_body(doc, "创新方案：", "以校园社区和服务入口为基础，引入 AI 对公开通知进行分类、摘要与结构化抽取，将信息转化为可订阅、可提醒、可追溯的校园事件。")
    add_labeled_body(doc, "应用价值：", "现有 Android 原型和服务端已覆盖真实校园场景；鸿蒙赛事版将把高频服务延伸到服务卡片、系统通知及跨设备协同，缩短从发现信息到完成行动的路径。")

    add_heading(doc, "三、真实效果与交互说明")
    add_screenshot_pair(doc)

    add_page_break(doc)
    add_heading(doc, "三、真实效果与交互说明（续）")
    add_single_screenshot(doc)
    add_labeled_body(doc, "统一入口：", "首页承载帖子浏览、互动与内容发现，校园页聚合竞赛通知和常用功能，侧边服务面板按任务场景组织入口。")
    add_labeled_body(doc, "智能处理：", "赛事版 AI 流程对通知原文执行清洗、分类、摘要和时间地点抽取，输出先经过规则校验与来源回链，再由用户确认。")
    add_labeled_body(doc, "闭环触达：", "确认后的事件可进入校历、服务卡片或系统通知；反馈与纠错返回服务端，持续优化标签、排序和提醒策略。")

    add_page_break(doc)
    add_heading(doc, "四、技术方案与 AI 处理流程")
    add_heading(doc, "（一）总体架构", level=2)
    architecture = doc.add_paragraph()
    architecture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    architecture.paragraph_format.space_after = Pt(6)
    r = architecture.add_run("ArkUI 页面与服务卡片  →  REST API 网关  →  Go 业务服务  →  Python 数据与 AI 服务  →  PostgreSQL")
    set_run_font(r, 10, bold=True, color=COLOR_ACCENT_DARK)

    add_heading(doc, "（二）AI 信息处理闭环", level=2)
    add_labeled_body(doc, "步骤 1 采集与清洗：", "从学校公开通知、竞赛资讯及用户授权内容中提取正文，保留来源链接和发布时间。")
    add_labeled_body(doc, "步骤 2 理解与抽取：", "大模型完成主题分类、摘要以及对象、时间、地点、截止日期等字段抽取，输出固定 JSON 结构。")
    add_labeled_body(doc, "步骤 3 校验与触达：", "规则引擎检查日期格式、重复项和来源可信度；低置信度结果进入人工确认，高置信度结果生成事件卡片与提醒。")

    add_heading(doc, "（三）实现基础与鸿蒙适配", level=2)
    add_technical_table(doc)
    add_labeled_body(doc, "状态说明：", "表中“已完成基础”均对应当前仓库与 Android 原型；ArkTS/ArkUI、服务卡片、系统通知和跨设备协同属于本次鸿蒙赛事版的实施计划。")

    add_page_break(doc)
    add_heading(doc, "五、作品介绍文档（800 字以内）")
    add_heading(doc, "（一）创意背景", level=2)
    intro_parts = [
        "高校信息分散在教务网站、学院通知、班级群和社区平台，学生需要反复查找，重要竞赛、考试与校园事务容易遗漏。沈理校园面向沈阳理工大学学生，以真实校园数据和社区需求为基础，建设信息聚合、校园互助和办事服务的一体化入口。",
        "作品已完成 Flutter/Android 原型以及 Go、Python、PostgreSQL 后端，支持校园社区信息流、通知与竞赛资讯、教务与成绩查询、校园集市、组队、评价及消息等功能。用户在首页浏览互动，在校园页查看聚合资讯与常用服务，在侧边面板快速进入竞赛、成绩、考试和反馈。赛事版 AI 服务将在既有数据链路上对学校公开通知和用户授权内容进行清洗、分类、摘要与字段抽取，生成对象、时间、地点、截止日期等可校验数据，经规则去重和人工确认后形成事件卡片，实现“发现信息—理解要点—加入日程—按时触达”的闭环。",
        "客户端当前采用 Flutter，业务接口为 REST+JSON，Go 服务负责鉴权与业务编排，Python 服务承载校园数据采集和智能处理，PostgreSQL 存储结构化数据，Nginx 与 Docker 用于部署。鸿蒙赛事版计划使用 ArkTS、ArkUI 与 Stage 模型重构首页、校园服务和消息提醒，复用现有服务端，并接入服务卡片、通知提醒及跨设备协同；敏感信息最小化上传，模型输出经过模式校验、来源回链和用户确认。",
        "创新点不是简单聚合入口，而是把 AI 信息理解、校园社区反馈和鸿蒙系统能力组合为可执行服务：公开信息按需求精准筛选，非标准通知转化为结构化事件，重要事项在卡片和通知中持续可见；社区互动反向补充需求与纠错。项目已有可运行原型、接口和真实场景数据，适合在校内社群、学院和竞赛组织中小范围验证，后续可按学校配置数据源和服务模块，形成可复制的高校智能服务方案。",
    ]
    headings = ["（一）创意背景", "（二）核心功能与交互", "（三）技术实现路径", "（四）创新点与应用前景"]
    # 第一项标题已添加，此处顺序写入四段并补齐后续小标题。
    add_body(doc, intro_parts[0])
    for heading, paragraph in zip(headings[1:], intro_parts[1:]):
        add_heading(doc, heading, level=2)
        add_body(doc, paragraph)

    intro_count = len("".join("".join(intro_parts).split()))
    if intro_count > 800:
        raise RuntimeError(f"作品介绍文档超出 800 字限制：{intro_count}")

    doc.core_properties.title = "沈理校园（SYLUlive）作品说明文档（初赛）"
    doc.core_properties.subject = "2026中国高校计算机大赛人工智能创意赛鸿蒙赛道"
    doc.core_properties.author = "咕咕嘎嘎"
    doc.core_properties.keywords = "鸿蒙, 人工智能, 校园服务, SYLUlive"
    doc.save(OUTPUT)
    restore_preserve_only_parts()
    print(f"OUTPUT={OUTPUT}")
    print(f"INTRO_COUNT={intro_count}")


if __name__ == "__main__":
    build()
